#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
FOCUST evaluation report generator.

Modes:
- basic: summary report from evaluation outputs (CSV + summary json)
- regenerated: dual-matching report (IoU vs center-distance) from successful_results_full.json
"""

import argparse
import csv
import json
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Sequence, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def read_json(path: Path) -> Any:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def _safe_int(v: Any, default: int = 0) -> int:
    try:
        if v is None:
            return default
        if isinstance(v, bool):
            return int(v)
        if isinstance(v, (int, float)):
            return int(v)
        return int(float(str(v)))
    except Exception:
        return default


def _safe_float(v: Any, default: float = 0.0) -> float:
    try:
        if v is None:
            return default
        if isinstance(v, bool):
            return float(int(v))
        if isinstance(v, (int, float)):
            return float(v)
        return float(str(v))
    except Exception:
        return default


def _prf(tp: int, fp: int, fn: int) -> Tuple[float, float, float]:
    precision = tp / (tp + fp) if (tp + fp) > 0 else 0.0
    recall = tp / (tp + fn) if (tp + fn) > 0 else 0.0
    f1 = 2 * precision * recall / (precision + recall) if (precision + recall) > 0 else 0.0
    return precision, recall, f1


def parse_iou_row(csv_path: Path, mode_key: str, iou_thr: float) -> Optional[Dict[str, Any]]:
    if not csv_path.exists():
        return None
    with csv_path.open("r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            try:
                mode = row.get("Mode") or row.get("模式") or row.get("mode")
                thr = float(row.get("IoU Threshold") or row.get("阈值") or row.get("IoU") or 0)
            except Exception:
                continue
            if (mode_key.lower() in (mode or "").lower()) and abs(thr - iou_thr) < 1e-6:
                return {
                    "Mode": mode,
                    "IoU": thr,
                    "Total GT": _safe_int(row.get("Total GT", 0)),
                    "Total Detections": _safe_int(row.get("Total Detections", 0)),
                    "TP": _safe_int(row.get("TP", 0)),
                    "FP": _safe_int(row.get("FP", 0)),
                    "FN": _safe_int(row.get("FN", 0)),
                    "Recall": _safe_float(row.get("Recall", 0.0)),
                    "Precision": _safe_float(row.get("Precision", 0.0)),
                    "F1 Score": _safe_float(row.get("F1 Score", 0.0)),
                }
    return None


def parse_final_stats(text_path: Path) -> Dict[str, Any]:
    stats: Dict[str, Any] = {}
    if not text_path.exists():
        return stats
    content = text_path.read_text(encoding="utf-8", errors="ignore")

    def find_num(label: str):
        m = re.search(label + r"\s*[:：]\s*([0-9\.]+)", content)
        return m.group(1) if m else None

    def find_int(label: str):
        v = find_num(label)
        return _safe_int(v, None) if v is not None else None

    def find_float(label: str):
        v = find_num(label)
        return _safe_float(v, None) if v is not None else None

    stats["total_sequences"] = find_int(r"总序列数")
    stats["success"] = find_int(r"成功处理")
    stats["failed"] = find_int(r"失败序列")
    stats["total_gt"] = find_int(r"总真值目标数")
    stats["total_dets"] = find_int(r"总检测目标数")
    stats["avg_time_per_seq"] = find_float(r"平均处理时间")
    return stats


# ------------------ Docx helpers ------------------

def _set_cn_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)

    for h in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        if h in doc.styles:
            s = doc.styles[h]
            s.font.name = "黑体"
            s._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def _set_run_font(run, *, size_pt: Optional[float] = None, bold: Optional[bool] = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold


def _set_paragraph_format(p) -> None:
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.25


def add_title(doc: Document, text: str, subtitle: Optional[str] = None) -> None:
    p = doc.add_paragraph()
    _set_paragraph_format(p)
    run = p.add_run(text)
    _set_run_font(run, size_pt=18, bold=True)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if subtitle:
        p2 = doc.add_paragraph()
        _set_paragraph_format(p2)
        run2 = p2.add_run(subtitle)
        _set_run_font(run2, size_pt=12, bold=False)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    size_map = {1: 14, 2: 13, 3: 12, 4: 11}
    p = doc.add_paragraph()
    _set_paragraph_format(p)
    run = p.add_run(text)
    _set_run_font(run, size_pt=size_map.get(level, 12), bold=True)


def add_paragraph(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    _set_paragraph_format(p)
    run = p.add_run(text)
    _set_run_font(run, size_pt=11, bold=bold)


def add_bullets(doc: Document, items: Sequence[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        _set_paragraph_format(p)
        run = p.add_run(item)
        _set_run_font(run, size_pt=11, bold=False)


def add_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[Any]],
    *,
    font_size_pt: float = 9.5,
    alignment: WD_TABLE_ALIGNMENT = WD_TABLE_ALIGNMENT.CENTER,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = alignment
    if "Table Grid" in doc.styles:
        table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.text = ""
        run = p.add_run(str(h))
        _set_run_font(run, size_pt=font_size_pt, bold=True)

    for row in rows:
        row_cells = table.add_row().cells
        for i, v in enumerate(row):
            p = row_cells[i].paragraphs[0]
            p.text = ""
            run = p.add_run(str(v))
            _set_run_font(run, size_pt=font_size_pt, bold=False)


def add_picture_if_exists(doc: Document, path: Path, width_inch: float = 6.0, caption: Optional[str] = None) -> None:
    if path and path.exists():
        doc.add_picture(str(path), width=Inches(width_inch))
        last_par = doc.paragraphs[-1]
        last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            cap = doc.add_paragraph(caption)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ------------------ Basic report ------------------

def build_basic_report(eval_dir: Path, output_path: Path, iou_key: float = 0.5) -> None:
    summary_json = read_json(eval_dir / "evaluation_summary.json")
    config_used = read_json(eval_dir / "config_used_for_evaluation.json")
    overall_csv = eval_dir / "evaluation_iou_sweep_report_overall.csv"
    with_filter_csv = eval_dir / "evaluation_iou_sweep_report_with_filter.csv"
    without_filter_csv = eval_dir / "evaluation_iou_sweep_report_without_filter.csv"
    final_stats = parse_final_stats(eval_dir / "evaluation_final_statistics.txt")
    dual_cmp_json = read_json(eval_dir / "dual_mode_comparison_data.json")

    overall_row = parse_iou_row(overall_csv, "overall", iou_key)
    with_row = parse_iou_row(with_filter_csv, "with", iou_key)
    without_row = parse_iou_row(without_filter_csv, "without", iou_key)

    with_prf = (dual_cmp_json.get("statistics") or {}).get("with_filter", {}) if isinstance(dual_cmp_json, dict) else {}
    without_prf = (dual_cmp_json.get("statistics") or {}).get("without_filter", {}) if isinstance(dual_cmp_json, dict) else {}

    doc = Document()
    _set_cn_styles(doc)
    add_title(doc, "Focust系统检测功能评估报告")

    add_heading(doc, "摘要", level=1)
    add_paragraph(
        doc,
        "本报告基于 FOCUST 检测评估输出生成，包含整体 IoU 统计与可视化结果。"
    )

    add_heading(doc, "评估概览", level=1)
    items = []
    seq_total = final_stats.get("total_sequences")
    seq_success = final_stats.get("success")
    seq_failed = final_stats.get("failed")
    if seq_total is None:
        seq_total = summary_json.get("sequence_total") if isinstance(summary_json, dict) else None
    if seq_success is None:
        seq_success = summary_json.get("sequence_success") if isinstance(summary_json, dict) else None
    if seq_failed is None:
        seq_failed = summary_json.get("sequence_failed") if isinstance(summary_json, dict) else None

    if seq_total is not None:
        items.append(f"总序列数: {seq_total}")
    if seq_success is not None:
        items.append(f"成功处理: {seq_success}")
    if seq_failed is not None:
        items.append(f"失败序列: {seq_failed}")
    if final_stats.get("total_gt") is not None:
        items.append(f"总真值目标数: {final_stats.get('total_gt')}")
    if final_stats.get("total_dets") is not None:
        items.append(f"总检测目标数: {final_stats.get('total_dets')}")

    if isinstance(config_used, dict):
        engine = config_used.get("engine")
        if engine:
            items.append(f"推理引擎: {engine}")
        models = config_used.get("models") if isinstance(config_used.get("models"), dict) else {}
        if isinstance(models, dict):
            if models.get("binary_classifier"):
                items.append(f"二分类权重: {models.get('binary_classifier')}")
            if models.get("multiclass_classifier"):
                items.append(f"多分类权重: {models.get('multiclass_classifier')}")
            if models.get("yolo_model"):
                items.append(f"YOLO权重: {models.get('yolo_model')}")
    if items:
        add_bullets(doc, items)

    add_heading(doc, f"IoU={iou_key:.2f} 指标对比", level=1)
    rows = []
    for row in (overall_row, with_row, without_row):
        if not row:
            continue
        rows.append([
            row.get("Mode", "-"),
            f"{row.get('IoU', 0):.3f}",
            row.get("Total GT", "-"),
            row.get("Total Detections", "-"),
            f"{row.get('Precision', 0.0):.4f}",
            f"{row.get('Recall', 0.0):.4f}",
            f"{row.get('F1 Score', 0.0):.4f}",
        ])
    if rows:
        add_table(doc, ["模式", "IoU", "总真值", "检测数", "Precision", "Recall", "F1"], rows)

    if with_prf or without_prf:
        add_heading(doc, "双模式宏平均指标", level=2)
        rows = []
        if with_prf:
            rows.append(["with_filter", f"{_safe_float(with_prf.get('precision')):.4f}", f"{_safe_float(with_prf.get('recall')):.4f}", f"{_safe_float(with_prf.get('f1_score')):.4f}"])
        if without_prf:
            rows.append(["without_filter", f"{_safe_float(without_prf.get('precision')):.4f}", f"{_safe_float(without_prf.get('recall')):.4f}", f"{_safe_float(without_prf.get('f1_score')):.4f}"])
        add_table(doc, ["模式", "Precision", "Recall", "F1"], rows)

    # 常用图表（如存在）
    add_heading(doc, "可视化图表", level=1)
    add_picture_if_exists(doc, eval_dir / "dual_mode_analysis" / "visualizations" / "dual_mode_performance_comparison.png", caption="双模式性能对比")
    add_picture_if_exists(doc, eval_dir / "dual_mode_analysis" / "visualizations" / "dual_mode_scatter_comparison.png", caption="双模式散点对比")
    add_picture_if_exists(doc, eval_dir / "dual_mode_with_filter" / "overall_performance.png", caption="启用过滤整体性能")
    add_picture_if_exists(doc, eval_dir / "dual_mode_without_filter" / "overall_performance.png", caption="禁用过滤整体性能")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


# ------------------ Regenerated report ------------------

def _find_eval_root(start_dir: Path) -> Path:
    cur = start_dir.resolve()
    for _ in range(30):
        if (cur / "config_used_for_evaluation.json").exists():
            return cur
        if cur.parent == cur:
            break
        cur = cur.parent
    raise FileNotFoundError("未找到 config_used_for_evaluation.json")


def _resolve_regenerated_summary(eval_dir: Path) -> Path:
    candidate = eval_dir / "evaluation_summary.json"
    if candidate.exists():
        return candidate
    summaries = sorted(eval_dir.glob("dataset_evaluation_*/evaluation_summary.json"))
    if summaries:
        return summaries[-1]
    raise FileNotFoundError("未找到 evaluation_summary.json")


def _aggregate_matching_metrics(results: List[Dict[str, Any]], mode_key: str) -> Dict[str, Any]:
    tp = fp = fn = 0
    sum_iou = 0.0
    sum_cd = 0.0
    matched = 0
    for row in results:
        if not isinstance(row, dict):
            continue
        if row.get("status") not in (None, "success"):
            continue

        metrics = None
        mbm = row.get("metrics_by_matching") or {}
        if isinstance(mbm, dict):
            bucket = mbm.get(mode_key) or mbm.get(mode_key.replace("_", "-"))
            if isinstance(bucket, dict):
                metrics = bucket.get("strict") or bucket.get("metrics")
        if not isinstance(metrics, dict):
            metrics = row.get("metrics") if isinstance(row.get("metrics"), dict) else None
        if not isinstance(metrics, dict):
            continue

        tp += _safe_int(metrics.get("tp"))
        fp += _safe_int(metrics.get("fp"))
        fn += _safe_int(metrics.get("fn"))
        if metrics.get("avg_iou") is not None:
            sum_iou += _safe_float(metrics.get("avg_iou")) * _safe_int(metrics.get("tp"))
        if metrics.get("avg_center_distance") is not None:
            sum_cd += _safe_float(metrics.get("avg_center_distance")) * _safe_int(metrics.get("tp"))
        matched += _safe_int(metrics.get("tp"))

    precision, recall, f1 = _prf(tp, fp, fn)
    return {
        "tp": tp,
        "fp": fp,
        "fn": fn,
        "precision": precision,
        "recall": recall,
        "f1": f1,
        "avg_iou": (sum_iou / matched) if matched else 0.0,
        "avg_center_distance": (sum_cd / matched) if matched else 0.0,
    }


def _aggregate_per_class(results: List[Dict[str, Any]], mode_key: str) -> Dict[str, Dict[str, Any]]:
    totals: Dict[str, Dict[str, int]] = {}
    for row in results:
        if not isinstance(row, dict):
            continue
        if row.get("status") not in (None, "success"):
            continue

        per_class = None
        mbm = row.get("metrics_by_matching") or {}
        if isinstance(mbm, dict):
            bucket = mbm.get(mode_key) or mbm.get(mode_key.replace("_", "-"))
            if isinstance(bucket, dict):
                per_class = bucket.get("per_class_strict")
        if not isinstance(per_class, dict):
            adv = row.get("advanced_results") or {}
            per_class = adv.get("per_class_strict") if isinstance(adv, dict) else None
        if not isinstance(per_class, dict):
            per_class = row.get("per_class") if isinstance(row.get("per_class"), dict) else None
        if not isinstance(per_class, dict):
            continue

        for cid, info in per_class.items():
            entry = totals.setdefault(str(cid), {"tp": 0, "fp": 0, "fn": 0, "gt": 0, "pred": 0})
            entry["tp"] += _safe_int(info.get("tp"))
            entry["fp"] += _safe_int(info.get("fp"))
            entry["fn"] += _safe_int(info.get("fn"))
            entry["gt"] += _safe_int(info.get("gt", info.get("gt_count", 0)))
            entry["pred"] += _safe_int(info.get("pred", info.get("det_count", 0)))

    out: Dict[str, Dict[str, Any]] = {}
    for cid, entry in totals.items():
        p, r, f1 = _prf(entry["tp"], entry["fp"], entry["fn"])
        out[cid] = {**entry, "precision": p, "recall": r, "f1": f1}
    return out


def build_regenerated_report(eval_dir: Path, output_path: Path) -> None:
    summary_path = _resolve_regenerated_summary(eval_dir)
    summary_json = read_json(summary_path)

    eval_root = _find_eval_root(summary_path.parent)
    results_path = eval_root / "successful_results_full.json"
    results = read_json(results_path)
    if not isinstance(results, list):
        results = []

    iou_metrics = _aggregate_matching_metrics(results, "iou")
    cd_metrics = _aggregate_matching_metrics(results, "center_distance")

    iou_per_class = _aggregate_per_class(results, "iou")
    cd_per_class = _aggregate_per_class(results, "center_distance")

    doc = Document()
    _set_cn_styles(doc)
    add_title(doc, "FOCUST 再生成评估报告", subtitle=str(summary_path.parent))

    add_heading(doc, "评估概览", level=1)
    add_bullets(doc, [
        f"评估根目录: {eval_root}",
        f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
    ])

    add_heading(doc, "匹配模式总体指标", level=1)
    rows = [
        ["IoU", f"{iou_metrics['precision']:.4f}", f"{iou_metrics['recall']:.4f}", f"{iou_metrics['f1']:.4f}", f"{iou_metrics['avg_iou']:.4f}", f"{iou_metrics['avg_center_distance']:.2f}"],
        ["中心距离", f"{cd_metrics['precision']:.4f}", f"{cd_metrics['recall']:.4f}", f"{cd_metrics['f1']:.4f}", f"{cd_metrics['avg_iou']:.4f}", f"{cd_metrics['avg_center_distance']:.2f}"],
    ]
    add_table(doc, ["匹配模式", "Precision", "Recall", "F1", "Avg IoU", "Avg CenterDist"], rows)

    if iou_per_class:
        add_heading(doc, "IoU 匹配 - 分类别指标", level=2)
        rows = []
        for cid, entry in sorted(iou_per_class.items(), key=lambda x: int(x[0])):
            rows.append([
                cid,
                entry.get("gt", 0),
                entry.get("pred", 0),
                f"{entry.get('precision', 0.0):.4f}",
                f"{entry.get('recall', 0.0):.4f}",
                f"{entry.get('f1', 0.0):.4f}",
            ])
        add_table(doc, ["Class", "GT", "Pred", "Precision", "Recall", "F1"], rows)

    if cd_per_class:
        add_heading(doc, "中心距离匹配 - 分类别指标", level=2)
        rows = []
        for cid, entry in sorted(cd_per_class.items(), key=lambda x: int(x[0])):
            rows.append([
                cid,
                entry.get("gt", 0),
                entry.get("pred", 0),
                f"{entry.get('precision', 0.0):.4f}",
                f"{entry.get('recall', 0.0):.4f}",
                f"{entry.get('f1', 0.0):.4f}",
            ])
        add_table(doc, ["Class", "GT", "Pred", "Precision", "Recall", "F1"], rows)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


def main() -> None:
    parser = argparse.ArgumentParser(description="FOCUST 评估报告生成器")
    parser.add_argument("--mode", choices=["basic", "regenerated"], default="basic", help="报告类型")
    parser.add_argument("--eval-dir", required=True, help="评估输出目录")
    parser.add_argument("--output", default="", help="输出 docx 路径")
    parser.add_argument("--iou", type=float, default=0.5, help="basic 模式使用的 IoU 阈值")

    args = parser.parse_args()
    eval_dir = Path(args.eval_dir)
    if not eval_dir.exists():
        raise FileNotFoundError(f"评估目录不存在: {eval_dir}")

    if args.mode == "basic":
        out_path = Path(args.output) if args.output else (eval_dir / "focust_report.docx")
        build_basic_report(eval_dir, out_path, iou_key=float(args.iou))
    else:
        out_path = Path(args.output) if args.output else (eval_dir / "regenerated_evaluation_report.docx")
        build_regenerated_report(eval_dir, out_path)


if __name__ == "__main__":
    main()
