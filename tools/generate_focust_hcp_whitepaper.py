#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Generate a Word (docx) whitepaper for the FOCUST system:
HCP algorithm -> binary screening -> multiclass re-screening -> temporal colony counting & classification.

This document is generated from the repository's current implementation and configs.
"""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REPO_ROOT = Path(__file__).resolve().parents[1]


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore") if path.exists() else ""


def _read_json(path: Path) -> Dict[str, Any]:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def set_cn_styles(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)

    for h in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        if h in doc.styles:
            s = doc.styles[h]
            s.font.name = "黑体"
            s._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_title(doc: Document, title: str, subtitle: Optional[str] = None) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if subtitle:
        p2 = doc.add_paragraph()
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(12)
        run2.font.name = "宋体"
        run2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_kv_table(
    doc: Document,
    rows: List[Tuple[str, str]],
    col_widths_inch: Tuple[float, float] = (2.2, 4.8),
) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in rows:
        r = table.add_row().cells
        r[0].text = str(k)
        r[1].text = str(v)
    for row in table.rows:
        row.cells[0].width = Inches(col_widths_inch[0])
        row.cells[1].width = Inches(col_widths_inch[1])


def add_param_table(
    doc: Document,
    title: str,
    params: Dict[str, Any],
    notes: Optional[Dict[str, str]] = None,
) -> None:
    doc.add_heading(title, level=2)

    items: List[Tuple[str, Any]] = []
    for k, v in params.items():
        if k.startswith("_"):
            continue
        items.append((k, v))

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    hdr[0].text = "参数"
    hdr[1].text = "值"
    hdr[2].text = "说明/来源"

    for k, v in sorted(items, key=lambda x: x[0]):
        row = table.add_row().cells
        row[0].text = str(k)
        row[1].text = json.dumps(v, ensure_ascii=False) if isinstance(v, (dict, list)) else str(v)
        row[2].text = (notes or {}).get(k, "")


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(9.5)


def add_picture_if_exists(doc: Document, img_path: Path, width_inch: float, caption: Optional[str] = None) -> None:
    if not img_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(img_path), width=Inches(width_inch))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _section_repo_basis(doc: Document) -> None:
    doc.add_heading("本文档依据与范围", level=1)
    doc.add_paragraph(
        "本文档以当前仓库实现为准，围绕 FOCUST 系统的“传统机器视觉预处理（HCP）→二分类初筛→多分类复筛”的三层架构，"
        "阐述从时序图像序列到最终菌落计数与菌种分类的端到端处理流程。文档内容尽量对应到真实代码接口、配置项与训练脚本，"
        "同时补充对含食品残渣场景下传统视觉与深度学习方法的优劣对比与融合策略论证。"
    )
    doc.add_paragraph("关键实现入口（可在仓库中直接打开）：")
    for path in [
        "detection/core/hpyer_core_processor.py",
        "detection/modules/enhanced_classification_manager.py",
        "detection/utils/classification_utils.py",
        "bi_train/train/classification_model.py",
        "mutil_train/train/classification_model.py",
        "bi_train/train/train_utils.py",
        "mutil_train/train/train_utils.py",
        "config/focust_config.json",
        "README.md",
    ]:
        add_code_block(doc, f"- {path}")


def _section_problem_background(doc: Document) -> None:
    doc.add_heading("问题背景与任务定义", level=1)
    doc.add_heading("应用场景：含食品残渣的时序菌落生长图像", level=2)
    doc.add_paragraph(
        "系统输入通常为同一培养皿在固定时间间隔下采集的时序图像序列（常见为约 40 帧）。在食品样本（复杂基质）环境中，"
        "培养基表面常出现大量颗粒状/斑点状的残渣、气泡、划痕与反光，外观与小菌落高度相似。与此同时，真实菌落具有“随时间持续生长/扩张”的动态特征。"
        "因此，单帧静态检测很容易产生高误检；而纯深度学习端到端在数据规模不足、基质变化强、跨批次光照漂移时，也会出现明显域偏移。"
    )
    doc.add_heading("系统目标：检测 + 计数 + 分类（两级分类）", level=2)
    doc.add_paragraph(
        "FOCUST 的核心目标可以拆解为三项：\n"
        "（1）在全帧尺度定位所有可能的菌落候选区域；\n"
        "（2）对候选区域进行“菌落/非菌落（残渣等）”二分类初筛；\n"
        "（3）对确认的菌落区域进行菌种（多类别）复筛，得到每个菌落的类别标签，并统计每类数量与总数。"
    )


def _section_system_overview(doc: Document, cfg: Dict[str, Any]) -> None:
    doc.add_heading("系统总体架构（HCP→二分类→多分类）", level=1)
    doc.add_paragraph(
        "系统采用“传统机器视觉 + 时序深度学习”的混合式三级流水线："
        "第一层 HCP 负责在全帧尺度快速提出候选目标（候选框集合）；"
        "第二层二分类网络在候选框层面去除大部分残渣/伪目标；"
        "第三层多分类网络在剩余目标上输出菌种类别。"
        "通过将计算量最大的深度学习推理限定在少量候选区域上，系统在保证精度的同时强调泛化能力与轻量化部署。"
    )

    doc.add_heading("配置入口与关键参数所在位置", level=2)
    doc.add_paragraph("系统配置集中在：")
    add_code_block(doc, "config/focust_config.json")
    doc.add_paragraph("与三层架构直接相关的配置块包括：")
    add_code_block(doc, "- hcp_algorithm：HCP 参数（背景帧数、阈值、最小面积、增长验证等）")
    add_code_block(doc, "- pretrained_models：二分类/多分类模型路径")
    add_code_block(doc, "- training_hyperparameters / model_architecture：训练与网络结构的默认参数")

    hcp_cfg = (cfg or {}).get("hcp_algorithm", {}) if isinstance(cfg, dict) else {}
    if hcp_cfg:
        add_param_table(
            doc,
            "HCP 关键参数（来自 config/focust_config.json:hcp_algorithm）",
            hcp_cfg,
            notes={
                "num_bg_frames": "用于背景建模的前 N 帧；对应 HCP 阶段1（背景中位数建模）",
                "bf_diameter": "双边滤波直径；对应 HCP 阶段0（预处理）",
                "min_colony_area_px": "最小菌落面积阈值（像素）；用于去除噪点/小伪目标",
                "bio_validation_enable": "是否开启生物学增长验证逻辑（随实现版本可能有所演进）",
                "min_growth_slope_threshold": "增长趋势阈值（用于判定是否存在持续增长）",
            },
        )


def _section_hcp_algorithm(doc: Document) -> None:
    doc.add_heading("第一层：HCP（HyperCoreProcessor）算法原理与实现细节", level=1)
    doc.add_paragraph(
        "HCP 是系统的传统机器视觉预处理层，核心思想是：在时序维度上做背景建模与变化检测，"
        "用“随时间产生的稳定变化”来区分真实菌落（持续生长）与静态/随机干扰（残渣、反光、噪声）。"
        "实现主体位于："
    )
    add_code_block(doc, "detection/core/hpyer_core_processor.py  （类：HpyerCoreProcessor）")

    doc.add_heading("输入/输出约定（工程接口）", level=2)
    doc.add_paragraph("输入：同一培养皿的帧路径列表 `image_paths`（按自然排序/时间顺序）。")
    doc.add_paragraph(
        "输出：`run()` 返回一个元组（用于 GUI/评估的兼容约定），其中最关键的是 `final_bboxes`："
        "每个 bbox 形如 `(x, y, w, h, id)`，其中 `id` 通常来自分割后的连通域标签。"
    )
    add_code_block(doc, "run() -> (frames_gray, binary_mask*255, time_color_map, final_labels, final_bboxes, ..., ...)")

    doc.add_heading("阶段0：并行加载与预处理", level=2)
    doc.add_paragraph(
        "HCP 首先用线程池并行读取灰度帧，并可选地对每帧执行双边滤波（bilateralFilter）以抑制噪声同时保留边缘。"
        "该步骤的目标是提升后续差分与阈值分割的稳定性。"
    )
    add_code_block(doc, "HpyerCoreProcessor._load_and_preprocess_frames()")

    doc.add_heading("阶段1：鲁棒背景建模 + 正/负信号解耦", level=2)
    doc.add_paragraph(
        "算法使用前 `num_bg_frames` 帧建立背景模型：对这些帧进行一致性修正（用锚帧替换瞬态差异像素），"
        "再取逐像素中位数作为背景 `background_model`，并计算噪声标准差图 `noise_model_std`。"
        "随后对后续数据帧做差分解耦：\n"
        "正信号（positive）= max(0, frame - background)\n"
        "负信号（negative）= max(0, background - frame)\n"
        "二者分别对应变亮/变暗的变化方向，用于适配不同培养基/光照条件下的菌落变化表现。"
    )
    add_code_block(doc, "HpyerCoreProcessor._stage1_model_decouple()")

    doc.add_heading("阶段1.5：时间一致性滤波（Temporal Consistency Filter）", level=2)
    doc.add_paragraph(
        "为了压制单帧随机噪声，HCP 引入时间一致性约束：仅当某像素在相邻两帧中都超过自适应阈值（由噪声标准差图×系数得到）时，"
        "才认为该像素属于稳定变化。该滤波特别适合区分“持续生长”的菌落与“偶发出现”的噪点/反光。"
    )
    add_code_block(doc, "HpyerCoreProcessor._apply_temporal_consistency_filter()")

    doc.add_heading("阶段2：鲁棒种子点追踪（关键帧 + 合并）", level=2)
    doc.add_paragraph(
        "HCP 会构建“首次出现时间图（first_appearance_map）”，并在若干关键帧上分析前景区域的连通性，"
        "追踪稳定生长的种子点。为了避免一个菌落产生多个相邻种子，算法还会对距离过近的种子点进行合并，并用出现时间更早者作为代表。"
    )
    add_code_block(doc, "HpyerCoreProcessor._stage2_find_robust_seeds()")
    add_code_block(doc, "HpyerCoreProcessor._track_stable_growth_seeds()  # key_frames = [0,2,5,7,12]")
    add_code_block(doc, "HpyerCoreProcessor._merge_close_seeds()")

    doc.add_heading("阶段3：掩码生成（FDT 阈值）+ 模糊菌落核提纯", level=2)
    doc.add_paragraph(
        "阶段3 通过最大强度投影（max over time）获得强变化图，再使用基于直方图形态的 First Drop Thresholding（FDT）"
        "得到初始二值掩码。随后针对“模糊菌落”（边缘梯度弱、边界不清晰），可选地执行核提纯："
        "利用梯度幅值图与自适应阈值（包含 Otsu 变体调整）提取更可靠的菌落核心区域，提升后续分割质量。"
    )
    add_code_block(doc, "threshold_first_drop(...)")
    add_code_block(doc, "HpyerCoreProcessor._stage3_generate_mask_and_refine_fuzzy(...)")

    doc.add_heading("阶段4：分水岭分割 + bbox 提取与尺寸过滤", level=2)
    doc.add_paragraph(
        "当存在鲁棒种子时，算法以 -max_intensity_map 作为“地形图”执行分水岭分割（watershed），"
        "以种子作为 marker 保持菌落实例的一致性；对于未被分水岭覆盖的孤立区域，会补充连通域标签。"
        "随后遍历 regionprops，输出每个实例的 bbox，并按尺寸阈值过滤：最小尺寸通常用实例像素面积（p.area）过滤噪点，"
        "最大尺寸可用 bbox 面积（p.bbox_area）过滤占据画面过大的伪目标。"
    )
    add_code_block(doc, "HpyerCoreProcessor._stage4_segment_and_extract_bboxes(...)")

    doc.add_heading("HCP 的优势与失效模式（工程视角）", level=2)
    doc.add_paragraph(
        "优势：无监督/少参数、速度快、可解释、对低数据量友好，并且利用时序变化这一“物理先验”天然区分静态残渣。"
        "常见失效模式：强光照漂移导致背景建模误差、培养基表面大范围变化导致掩码过大、极早期微小菌落变化弱导致漏检、"
        "以及种子点/分水岭对噪声敏感时带来的过分割/欠分割。"
        "这些失效模式正是第二层/第三层网络要补强的部分。"
    )


def _section_binary_model(doc: Document) -> None:
    doc.add_heading("第二层：二分类初筛网络（菌落/残渣）", level=1)
    doc.add_paragraph(
        "二分类网络的任务是：在 HCP 给出的候选框集合上，识别哪些候选为真实菌落，哪些为食品残渣、气泡、反光等非菌落目标。"
        "与单帧 CNN 不同，本系统明确利用时序信息：菌落具有稳定的增长纹理与形态演化，而残渣多为静态或随机变化。"
    )
    doc.add_paragraph("实现位置：")
    add_code_block(doc, "bi_train/train/classification_model.py  （类：Focust，特征提取器：BioGrowthNetV2）")

    doc.add_heading("输入输出与推理约定", level=2)
    doc.add_paragraph(
        "输入张量形状为 `[B, T, 3, H, W]`（默认 H=W=224）。推理端会对每个 bbox 在每帧上裁剪 ROI patch，"
        "并按序堆叠为时序序列。二分类输出为 2 类 logits。工程管线中约定：预测索引 0 代表“菌落”。"
    )
    add_code_block(doc, "detection/modules/enhanced_classification_manager.py: colony_class_index = 0")

    doc.add_heading("网络结构（按实现拆解）", level=2)
    doc.add_paragraph("该网络由三部分组成：时序特征提取（BioGrowthNetV2）→双路径时序建模（CfC/NCP）→跨注意力融合与分类器。")

    doc.add_heading("（1）BioGrowthNetV2：轻量级空间特征提取", level=3)
    doc.add_paragraph(
        "BioGrowthNetV2 采用 Depthwise Separable Conv + Inverted Residual（MBConv）风格，"
        "逐步下采样并提升通道数，最终通过全局池化与全连接层输出 feature_dim 维特征向量。"
        "该设计在保持表达能力的同时显著降低参数量，适合在候选框数量较多的场景做批量推理。"
    )

    doc.add_heading("（2）双路径 CfC/NCP：显式时序建模", level=3)
    doc.add_paragraph(
        "网络内部构建两条时序路径：\n"
        "Path1：以最后一帧特征为“终态权重”，对整个序列特征进行加权（强调最终形态）；\n"
        "Path2：对序列做全局时序注意力聚合，再将其扩展回序列维度（强调全局动态）。\n"
        "两条路径分别进入 CfCWrapper（基于 AutoNCP wiring 的 CfC 模块），输出时序摘要向量。"
    )
    add_code_block(doc, "bi_train/train/ncps/wirings.py  （AutoNCP）")
    add_code_block(doc, "bi_train/train/ncps/wrapper.py   （CfCWrapper）")

    doc.add_heading("（3）Cross-Attention Fusion + Classifier：融合两条时序信息", level=3)
    doc.add_paragraph(
        "两条路径的输出通过跨注意力融合模块进行融合，最后经过 MLP 分类器输出二分类结果。"
        "损失函数支持 CE / Focal / MSE，并在 auto 模式下根据类别不平衡程度自动切换到 FocalLoss。"
    )

    doc.add_heading("训练流程要点（bi_train）", level=2)
    doc.add_paragraph("训练入口脚本与核心循环：")
    add_code_block(doc, "bi_train/bi_training.py")
    doc.add_paragraph(
        "训练数据以 COCO 风格 annotations 组织，并按 sequence_id 将同一菌落的多帧 patch 组成序列。"
        "训练增强使用常规随机翻转、旋转、ColorJitter，并对二分类使用 ImageNet 归一化以保持与推理一致。"
        "训练过程包含 EMA 权重平滑、Early Stopping，并支持 Optuna 进行超参数搜索（见后文）。"
    )


def _section_multiclass_model(doc: Document, cfg: Dict[str, Any]) -> None:
    doc.add_heading("第三层：多分类复筛网络（菌种识别）", level=1)
    doc.add_paragraph(
        "多分类网络的任务是：在二分类确认的菌落 ROI 上，输出具体菌种类别（示例配置为 5 类）。"
        "与二分类相比，多分类更关注形态学差异与细粒度纹理，同时仍利用时序信息提升稳定性。"
    )
    doc.add_paragraph("实现位置：")
    add_code_block(doc, "mutil_train/train/classification_model.py  （类：Focust，特征提取器：SimpleCNNFeatureExtractor）")

    doc.add_heading("输入输出与类别映射", level=2)
    doc.add_paragraph(
        "输入同样为 `[B, T, 3, 224, 224]` 的 ROI 序列。多分类输出为 K 类 logits。"
        "由于数据集内部类别 id 可能并非从 0 连续编号，推理端通过 `multiclass_index_to_category_id_map` "
        "将模型输出索引映射到真实 category_id。"
    )
    add_code_block(doc, "detection/modules/enhanced_classification_manager.py: multiclass_id_map")
    if isinstance(cfg, dict):
        labels = cfg.get("class_labels", {}).get("zh_cn") or {}
        if labels:
            doc.add_paragraph("示例类别定义（来自 config/focust_config.json:class_labels.zh_cn）：")
            add_param_table(doc, "菌种类别（示例）", labels)

    doc.add_heading("网络结构（按实现拆解）", level=2)
    doc.add_paragraph(
        "多分类网络由：SimpleCNNFeatureExtractor（深度可分离卷积）→ 双路径时序建模（CfC + Conv1d）→ EnhancedAttentionFusion → 分类层 组成。"
        "同时支持 data_mode='enhanced'：额外输入第二组序列 x2，用于增强特征融合（在数据具备双模态/增强帧时启用）。"
    )

    doc.add_heading("训练流程要点（mutil_train）", level=2)
    doc.add_paragraph("训练入口脚本与核心循环：")
    add_code_block(doc, "mutil_train/mutil_training.py")
    doc.add_paragraph(
        "训练增强包含随机裁剪（RandomResizedCrop）等操作；推理端采用 Resize(+CenterCrop) 保持尺度稳定。"
        "多分类训练脚本内包含显存估计与自动 batch_size 调整逻辑，避免 OOM。"
        "与二分类不同，多分类推理默认不做 ImageNet 归一化（与训练保持一致）。"
    )


def _section_hpo_optuna(doc: Document) -> None:
    doc.add_heading("Optuna 超参数搜索（HPO）与参数配置", level=1)

    doc.add_heading("二分类 HPO（bi_train）", level=2)
    doc.add_paragraph(
        "二分类采用 MedianPruner 进行剪枝，并将 study 持久化到 sqlite 文件（便于断点续跑）。"
        "目标函数在固定网络结构下，仅搜索少量关键超参。"
    )
    add_code_block(doc, "bi_train/train/train_utils.py: hyperparameter_optimization(), objective()")
    doc.add_paragraph("搜索空间（实现中的 trial.suggest_*）：")
    add_kv_table(
        doc,
        [
            ("lr", "log-uniform, 1e-5 ~ 1e-3"),
            ("weight_decay", "log-uniform, 1e-5 ~ 1e-3"),
            ("optimizer", "AdamW | RMSprop"),
            ("dropout_rate", "uniform, 0.2 ~ 0.5"),
        ],
    )
    doc.add_paragraph(
        "HPO 训练轮数较短（含 warmup），以验证集准确率为目标；若出现 CUDA OOM，会将该 trial 视为剪枝。"
        "结束后会保存 best_hpo_params.json、study 数据库以及可视化曲线（若环境支持）。"
    )

    doc.add_heading("多分类 HPO（mutil_train）", level=2)
    doc.add_paragraph(
        "多分类采用 TPE 采样 + MedianPruner。除学习率、权重衰减等外，还会搜索优化器的内部超参（Adam betas / SGD momentum / RMSprop alpha 等），"
        "并在必要时对 fusion_output_size 做安全修正，避免结构不合法。"
    )
    add_code_block(doc, "mutil_train/train/train_utils.py: hyperparameter_optimization(), objective()")
    doc.add_paragraph("搜索空间（核心项）：")
    add_kv_table(
        doc,
        [
            ("lr", "log-uniform, 1e-5 ~ 1e-2"),
            ("weight_decay", "log-uniform, 1e-6 ~ 1e-3"),
            ("optimizer", "Adam | SGD | RMSprop"),
            ("dropout_rate", "uniform, 0.1 ~ 0.5"),
            ("Adam.beta1", "uniform, 0.85 ~ 0.99"),
            ("Adam.beta2", "uniform, 0.9 ~ 0.9999"),
            ("Adam.epsilon", "log-uniform, 1e-8 ~ 1e-4"),
            ("SGD.momentum", "uniform, 0.0 ~ 0.99"),
            ("SGD.nesterov", "categorical, True/False"),
            ("RMSprop.momentum", "uniform, 0.0 ~ 0.99"),
            ("RMSprop.alpha", "uniform, 0.9 ~ 0.99"),
            ("RMSprop.epsilon", "log-uniform, 1e-8 ~ 1e-4"),
        ],
    )


def _section_inference_and_counting(doc: Document, cfg: Dict[str, Any]) -> None:
    doc.add_heading("端到端推理：从时序序列到菌落计数与分类", level=1)

    doc.add_heading("推理流水线总览", level=2)
    doc.add_paragraph(
        "推理阶段以一条序列（同一培养皿的一组帧路径）为单位。系统首先通过 HCP 获得候选框集合，"
        "再对每个候选框裁剪 ROI 序列送入二分类网络进行初筛，保留预测为“菌落”的候选，"
        "最后将这些候选送入多分类网络输出菌种类别，并汇总得到每类数量与总数。"
    )
    doc.add_paragraph("GUI/评估管线中的关键日志描述：")
    add_code_block(doc, "Pipeline: HCP detection -> binary screening -> multiclass -> IoU/class matching -> metrics")

    doc.add_heading("（1）候选框的 ROI 过滤：椭圆 ROI 与边缘忽略", level=2)
    doc.add_paragraph(
        "在培养皿图像中，边缘区域常出现光照反射与背景结构，误检概率高。系统提供椭圆 ROI 掩码（ellipse.png）"
        "用于过滤位于边缘的候选框，从而降低误检并提升稳定性。"
    )
    add_code_block(doc, "detection/modules/roi_utils.py  + ellipse.png")
    add_picture_if_exists(doc, REPO_ROOT / "ellipse.png", width_inch=2.2, caption="椭圆 ROI 掩码（示意）")

    doc.add_heading("（2）小菌落策略：标记为生长中并跳过多分类（可选）", level=2)
    doc.add_paragraph(
        "对于尺寸过小的目标，时序与纹理信息不足，多分类容易不稳定。系统提供“小菌落过滤/标记”策略："
        "当 bbox 尺寸小于阈值时，可将其标记为特殊类别（例如 0=未分类/生长中），并跳过多分类推理。"
        "该策略在质控场景下可显著降低误分类带来的下游干扰。"
    )

    doc.add_heading("（3）ROI 序列构建：SequenceDataManager（一次遍历 + 批量裁剪缓存）", level=2)
    doc.add_paragraph(
        "为了高效地对多个 bbox 构建时序序列，系统实现了 SequenceDataManager：只遍历一次所有帧，"
        "对每帧裁剪出所有 bbox 对应的 patch，并缓存到内存字典中；之后再为每个 bbox 组装出 `[T,3,224,224]` 的序列张量。"
        "该策略同时降低了磁盘重复读取与 Python 层循环开销。"
    )
    add_code_block(doc, "detection/utils/classification_utils.py: SequenceDataManager")

    doc.add_heading("（4）二分类初筛：过滤保留“菌落”候选", level=2)
    doc.add_paragraph(
        "二分类管理器（EnhancedClassificationManager）加载二分类模型后，对每个 bbox 的时序序列进行推理，"
        "并保留预测索引为 0 的 bbox。二分类 transform 使用 Resize + ImageNet Normalize（与训练一致）。"
    )
    add_code_block(doc, "detection/modules/enhanced_classification_manager.py: run_binary_classification()")

    doc.add_heading("（5）多分类复筛：输出菌种类别并执行索引到 category_id 映射", level=2)
    doc.add_paragraph(
        "多分类推理同样对 bbox 的时序序列进行预测，输出 logits 的 argmax 作为“预测索引”。"
        "由于真实类别 id 与输出索引不一定一致，系统通过 `multiclass_index_to_category_id_map` 完成映射，得到最终 category_id。"
    )
    add_code_block(doc, "detection/modules/enhanced_classification_manager.py: run_multiclass_classification()")
    mapping = (cfg or {}).get("models", {}).get("multiclass_index_to_category_id_map") if isinstance(cfg, dict) else None
    if isinstance(mapping, dict) and mapping:
        add_param_table(doc, "多分类索引到类别ID映射（来自配置）", mapping)

    doc.add_heading("（6）菌落计数与结果汇总", level=2)
    doc.add_paragraph(
        "最终菌落计数通常以“最终保留的 bbox 数量”为基础：总数 = len(final_bboxes)。"
        "分类计数则以 bbox→category_id 的映射结果做分组统计：对每个 category_id 计数即可得到每类菌落数量。"
        "在评估场景下，系统还会进一步与真值框匹配（IoU 或中心距离），统计 TP/FP/FN 与 Precision/Recall/F1。"
    )
    doc.add_paragraph("可复用的伪代码如下：")
    add_code_block(
        doc,
        "counts = {}\n"
        "for bbox_xywh in final_bboxes:\n"
        "    key = tuple(bbox_xywh[:4])\n"
        "    cid = predictions.get(key, -1)\n"
        "    counts[cid] = counts.get(cid, 0) + 1\n"
        "total = sum(counts.values())",
    )


def _section_comparison_and_hybrid_rationale(doc: Document) -> None:
    doc.add_heading("关键问题：传统机器视觉 vs 深度学习，以及为何采用融合方案", level=1)

    doc.add_heading("传统机器视觉（以 HCP 为代表）的优势", level=2)
    doc.add_paragraph(
        "（1）可解释性强：背景建模、阈值分割、连通域/分水岭等步骤清晰可追溯；\n"
        "（2）对数据依赖弱：无需大规模标注数据即可运行；\n"
        "（3）速度快、资源占用小：适合在边缘端快速提出候选；\n"
        "（4）可利用“时序变化”这种物理先验：对静态残渣具有天然抑制能力。"
    )

    doc.add_heading("传统机器视觉的弊端（含食品残渣场景）", level=2)
    doc.add_paragraph(
        "（1）阈值/参数敏感：光照漂移、培养基纹理变化会显著影响阈值分割与掩码质量；\n"
        "（2）复杂基质下易产生非典型伪目标：残渣可能随水分扩散产生缓慢变化，误导变化检测；\n"
        "（3）实例分割困难：菌落粘连、重叠时，连通域与分水岭可能过分割或欠分割；\n"
        "（4）跨场景泛化有限：不同培养基/相机/光学系统下需要重新调参。"
    )

    doc.add_heading("深度学习（时序网络）的优势", level=2)
    doc.add_paragraph(
        "（1）表征能力强：可自动学习复杂纹理与形态差异，适合细粒度菌种识别；\n"
        "（2）对复杂干扰更鲁棒：在充分数据覆盖下可吸收光照、残渣、噪声等变化；\n"
        "（3）可端到端优化：通过损失函数与数据增强提升整体任务指标；\n"
        "（4）时序建模可利用增长动力学：菌落的增长轨迹比静态外观更稳定、更具判别力。"
    )

    doc.add_heading("深度学习的弊端（含食品残渣场景）", level=2)
    doc.add_paragraph(
        "（1）数据依赖强：需要大量、覆盖多基质/多批次/多光照的标注样本，否则会出现域偏移；\n"
        "（2）计算资源与部署成本：全帧端到端检测/分割模型较重，实时性与边缘部署困难；\n"
        "（3）可解释性较弱：在误检/误分类时难以快速定位原因；\n"
        "（4）训练与维护成本高：需要持续的再训练、调参与版本管理。"
    )

    doc.add_heading("为何采用“传统机器视觉 + 深度学习”融合", level=2)
    doc.add_paragraph(
        "融合方案的核心动机是“互补与分工”：\n"
        "（1）泛化：HCP 利用时序变化先验做候选提取，可在数据不足时仍保持可用；深度模型专注在候选 ROI 上做判别，降低域偏移带来的灾难性误检；\n"
        "（2）轻量：HCP 将全帧问题降维为少量 ROI 分类，使二/多分类网络可以采用更轻的 backbone 与较小输入；\n"
        "（3）鲁棒：残渣的静态性/随机性会在 HCP 与二分类两层被逐步过滤；菌种差异则由多分类负责；\n"
        "（4）工程可控：每一层都有清晰指标与调参入口，可分层定位问题、分层迭代。"
    )


def _section_appendix_configs(doc: Document, cfg: Dict[str, Any]) -> None:
    doc.add_heading("附录：配置参数与训练/推理一致性检查点", level=1)
    if not isinstance(cfg, dict) or not cfg:
        doc.add_paragraph("未找到配置文件，跳过参数附录。")
        return

    train_hp = cfg.get("training_hyperparameters", {})
    model_arch = cfg.get("model_architecture", {})
    pretrained = cfg.get("pretrained_models", {})

    if isinstance(pretrained, dict):
        add_param_table(doc, "预训练模型路径（pretrained_models）", pretrained)
    if isinstance(train_hp, dict):
        add_param_table(doc, "训练超参（training_hyperparameters.binary）", train_hp.get("binary", {}))
        add_param_table(doc, "训练超参（training_hyperparameters.multiclass）", train_hp.get("multiclass", {}))
    if isinstance(model_arch, dict):
        add_param_table(doc, "模型结构（model_architecture.common）", model_arch.get("common", {}))
        add_param_table(doc, "模型结构（model_architecture.binary）", model_arch.get("binary", {}))
        add_param_table(doc, "模型结构（model_architecture.multiclass）", model_arch.get("multiclass", {}))

    doc.add_heading("训练/推理 transform 一致性（重要）", level=2)
    doc.add_paragraph(
        "二分类：训练与推理均使用 ImageNet Normalize；多分类：训练与推理均不使用 Normalize。"
        "该差异在代码中被明确固化（以与历史训练权重保持一致），修改时需同步调整训练与推理两端。"
    )
    add_code_block(doc, "detection/modules/enhanced_classification_manager.py: transforms for binary vs multiclass")
    add_code_block(doc, "bi_train/bi_training.py: train_transform includes Normalize")
    add_code_block(doc, "mutil_train/mutil_training.py: train_transform excludes Normalize")


def _section_limitations_and_open_issues(doc: Document) -> None:
    doc.add_heading("存在问题、不足与改进方向", level=1)

    doc.add_heading("算法层面的不足（HCP 与三层协同）", level=2)
    doc.add_paragraph(
        "（1）HCP 的参数敏感性：背景帧数、噪声阈值系数、FDT 阈值等对光照漂移与培养基纹理变化敏感；不同批次/相机/曝光条件下可能需要重新校准。\n"
        "（2）复杂基质的“慢变化伪目标”：食品残渣并非完全静态，受吸水扩散/溶解/反光变化影响可能产生缓慢变化，导致 HCP 将其作为候选。\n"
        "（3）实例分割不确定性：菌落粘连/重叠时，分水岭的 marker 依赖鲁棒种子，种子不稳会导致过分割或欠分割，从而影响计数。\n"
        "（4）极早期小菌落：信号弱、面积小，HCP 与二分类都可能漏检；多分类在该阶段更容易不稳定，因此工程上往往需要“小菌落标记为生长中/跳过多分类”的策略配合。"
    )

    doc.add_heading("深度学习层面的不足（二分类/多分类）", level=2)
    doc.add_paragraph(
        "（1）数据依赖与域偏移：即使使用时序网络，跨食品基质、跨培养基、跨光照条件仍可能出现分布漂移；需要持续补充覆盖性数据与再训练。\n"
        "（2）标签噪声的放大：如果上游候选框预标注存在系统性偏差（漏标/框偏/类别错），二分类与多分类会学习到偏差并在推理中放大。\n"
        "（3）推理一致性约束：二分类与多分类在 transform 上存在刻意差异（是否 Normalize），若训练/推理不一致会造成性能骤降；修改任何一端都需联动检查。"
    )

    doc.add_heading("工程与可维护性不足（当前代码实现可见的风险点）", level=2)
    doc.add_paragraph(
        "（1）模块命名与接口兼容：仓库中存在 Hpyer/Hyper 等拼写差异，部分工具脚本可能出现引用不一致并触发降级逻辑，影响批处理稳定性。\n"
        "（2）依赖可选导致能力降级：HCP 中部分功能依赖 `scipy`/`scikit-image`/`matplotlib`，缺失时会关闭高级能力（如 watershed/FDT 的部分实现与可视化）。\n"
        "（3）数据组织约束较强：例如 GUI 构建检测数据集默认只处理带 `_back` 的帧；若数据命名不符合约定会被跳过。\n"
        "（4）训练脚本的导入路径：训练模块在不同运行方式下可能触发相对/绝对导入差异，建议统一通过项目根目录运行，并保证包路径一致。"
    )

    doc.add_heading("可落地的改进方向（建议按收益优先级）", level=2)
    doc.add_paragraph(
        "（1）数据闭环优先：将“预标注→人工矫正→再训练→再预标注”做成固定迭代周期，并沉淀误检/漏检的困难样本集（hard cases）。\n"
        "（2）HCP 的自适应增强：对背景建模引入更强的光照归一/局部对比度稳定策略，降低不同批次参数漂移；同时对种子点策略做稳健化。\n"
        "（3）模型的不确定性输出：在二/多分类中暴露置信度或温度校准后的概率，给人工矫正与质控流程提供优先级排序。\n"
        "（4）一致性与可复现：将训练配置、模型 init_args、transform pipeline 与版本号固化输出到 checkpoint，并提供“一键复现”脚本。"
    )


def _section_binary_dataset_construction(doc: Document) -> None:
    doc.add_heading("二分类数据集构建（从检测数据集到菌落/非菌落）", level=1)
    doc.add_paragraph(
        "二分类数据集的核心目的是把“候选区域”转化为稳定的监督信号：正样本为真实菌落序列，负样本为非菌落（食品残渣、气泡、噪点等）序列。"
        "在工程实现中，二分类数据集通常从已有的检测数据集（含 sequence_id 与 bbox 标注）派生，"
        "通过类别映射、质量过滤与样本均衡，得到仅含 2 个类别的训练集。"
    )
    doc.add_paragraph("工具入口（GUI）：")
    add_code_block(doc, "gui/binary_dataset_builder.py  （BinaryDatasetBuilder / BinaryDatasetBuilderGUI）")

    doc.add_heading("输入数据格式（源检测数据集）", level=2)
    doc.add_paragraph(
        "源数据集需包含：`images/`、`annotations/annotations.json`。其中 `images` 条目必须带 `sequence_id` 与 `time`，"
        "annotations 条目带 `sequence_id`、`bbox`、`category_id`。该格式与系统内部的 SeqAnno/COCO 兼容写法一致。"
    )

    doc.add_heading("关键步骤 1：选择正/负类别（或负样本自动选择）", level=2)
    doc.add_paragraph(
        "构建时需要先指定“正样本类别”（positive_categories）。通常对应你认为是“真实菌落”的类别集合。"
        "负样本类别（negative_categories）可选：\n"
        "若显式指定，则只有命中这些类别的序列会作为负样本；\n"
        "若不指定，则采用自动模式：所有不属于正样本类别的序列都会被当作负样本候选。"
    )
    add_code_block(doc, "BinaryDatasetBuilder.set_positive_categories() / set_negative_categories()")

    doc.add_heading("关键步骤 2：按序列聚合并初分配正/负", level=2)
    doc.add_paragraph(
        "构建器先按 `sequence_id` 聚合 images 与 annotations，然后对每个序列检查其 annotation 的 category_id 集合：\n"
        "若命中任一正类别 → 判为正样本序列；\n"
        "否则 → 判为负样本序列（或在显式负类别模式下要求命中负类别）。"
    )
    add_code_block(doc, "BinaryDatasetBuilder._group_by_sequence() + build_binary_dataset() 初始分类逻辑")

    doc.add_heading("关键步骤 3：HCP 质量过滤（可选，但强烈建议开启）", level=2)
    doc.add_paragraph(
        "为了减少无效序列（帧数过少、无明显变化、图像异常）对训练的干扰，构建器提供 HCP 质量过滤："
        "对每条序列（通常取前若干帧）运行 HCP，并根据检测到的有效区域数量与总面积计算 quality_score，"
        "低于阈值的序列会被剔除；当 HCP 不可用或运行失败时会降级到基础质量检查。"
    )
    add_code_block(doc, "BinaryDatasetBuilder.apply_hcp_filtering() / _basic_quality_filter()")

    doc.add_heading("关键步骤 4：样本均衡（正负比）", level=2)
    doc.add_paragraph(
        "二分类常见问题是负样本远多于正样本。构建器支持按 `balance_ratio` 下采样负样本，"
        "目标为：negative_count ≈ positive_count × balance_ratio。"
    )
    add_code_block(doc, "BinaryDatasetBuilder.balance_samples()")

    doc.add_heading("关键步骤 5：导出为二分类数据集（2 类 + 可追溯字段）", level=2)
    doc.add_paragraph(
        "导出时会创建新的数据集结构，并将类别折叠为两类：id=1 positive，id=0 negative。"
        "同时会保留可追溯字段（例如 original_category_id、original_sequence_id），用于后续排查标签来源与错误。"
        "若源数据集存在 images2（增强帧）也会同步复制。"
    )
    add_code_block(doc, "BinaryDatasetBuilder._copy_sequence_data(): original_category_id/original_sequence_id")

    doc.add_heading("输出结构（典型）", level=2)
    add_code_block(
        doc,
        "output_binary_dataset/\n"
        "  images/{seq_id}/{seq_id}_{time}.jpg\n"
        "  (optional) images2/{seq_id}/...\n"
        "  annotations/annotations.json\n"
        "  build_info.json\n"
        "  dataset_stats.json",
    )


def _section_multiclass_dataset_prelabel_and_correction(doc: Document) -> None:
    doc.add_heading("多分类数据集构建：HCP+二分类预标注 → 人工矫正 → 训练用分类数据集导出", level=1)
    doc.add_paragraph(
        "多分类数据集的标注成本最高：既要有菌落定位框，又要有菌种类别。FOCUST 的推荐做法是利用 HCP+二分类先做“预标注框”，"
        "再由人工在可视化编辑器中快速矫正（删错框、补漏框、调位置、修类别），形成高质量真值。"
        "该闭环在含食品残渣场景尤其重要：机器预标注负责覆盖与速度，人工矫正负责精度与纠错。"
    )

    doc.add_heading("步骤 A：按菌种组织原始时序数据（决定“类别真值”）", level=2)
    doc.add_paragraph(
        "将原始时序序列按菌种/培养方法放入不同文件夹（folder label）。在预标注阶段，系统会把“文件夹所属菌种”作为该序列的类别真值来源，"
        "并将过滤后的菌落框统一标记为该类别。该假设要求：每个输入文件夹内的序列应尽量单一菌种，避免混入造成标签噪声。"
    )

    doc.add_heading("步骤 B：HCP 候选检测 + 二分类筛除残渣（生成预标注框）", level=2)
    doc.add_paragraph(
        "GUI 的“数据集构建”会启动后台 DetectionThread：\n"
        "（1）读取每个序列的帧路径（默认只取 `_back` 帧），并用 natsort 做时间排序；\n"
        "（2）运行 HpyerCoreProcessor 得到 initial_bboxes；\n"
        "（3）若配置了二分类模型 bi_cat98.pth，则用 EnhancedClassificationManager 过滤，得到 final_bboxes；"
        "否则退化为仅使用 HCP 结果。\n"
        "（4）把序列帧复制到输出数据集目录，并把 final_bboxes 写入 annotations.json，类别 id 对应当前菌种（来自 folder）。"
    )
    add_code_block(doc, "gui/threads.py: DetectionThread.run()  # Step 1/3~3/3")
    add_code_block(doc, "detection/core/hpyer_core_processor.py")
    add_code_block(doc, "detection/modules/enhanced_classification_manager.py")

    doc.add_heading("步骤 C：人工矫正（把“预标注”变成“真值”）", level=2)
    doc.add_paragraph(
        "预标注完成后，使用标注编辑器逐序列检查：\n"
        "删除明显残渣框（误检）、补充漏检菌落框（漏检）、调整框位置/大小（框偏）、修正类别（错分/混入），"
        "并保存更新后的 annotations.json。此阶段决定最终数据质量，也是提升泛化能力的关键。"
    )
    add_code_block(doc, "annotation_editor.py  （或 gui/annotation_editor.py 中的编辑器入口）")

    doc.add_heading("步骤 D：从检测数据集导出“多分类训练用分类数据集”（ROI 序列裁剪）", level=2)
    doc.add_paragraph(
        "当检测数据集的 bbox 与类别已被人工矫正后，可将其转换为分类训练数据集："
        "将每个 bbox 在每帧上裁剪为 224×224 patch，并以 sequence 形式组织为 `[T,3,224,224]` 的样本。"
        "该导出流程由 AnnotationEditor 的导出函数完成，并在 GUI 中由 ClassificationDatasetBuildThread 调用。"
    )
    add_code_block(doc, "gui/annotation_editor.py: generate_classification_dataset_with_messages()")
    add_code_block(doc, "gui/threads.py: ClassificationDatasetBuildThread")

    doc.add_heading("质量控制建议（强烈建议做）", level=2)
    doc.add_paragraph(
        "（1）对预标注结果按序列抽样质检：统计每序列框数量分布，快速发现 HCP 参数异常或二分类过严/过松。\n"
        "（2）优先矫正“模型不确定/小菌落/边缘区域”样本：这些样本对泛化最敏感。\n"
        "（3）保持类别一致：多分类最终类别需与配置中的 category_id/label map 对齐，避免训练/推理映射错位。"
    )


def main() -> Path:
    cfg_path = REPO_ROOT / "config" / "focust_config.json"
    cfg = _read_json(cfg_path)

    doc = Document()
    set_cn_styles(doc)

    add_picture_if_exists(doc, REPO_ROOT / "logo.png", width_inch=0.8)
    add_title(
        doc,
        "FOCUST 系统：HCP 算法处理—二分类初筛—多分类复筛\n时序菌落计数与分类技术白皮书",
        subtitle=f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
    )

    doc.add_paragraph()
    add_kv_table(
        doc,
        [
            ("仓库根目录", str(REPO_ROOT)),
            ("核心检测模块", "detection/core/hpyer_core_processor.py (HpyerCoreProcessor)"),
            ("二分类模型", "bi_train/train/classification_model.py (Focust/BioGrowthNetV2 + CfC)"),
            ("多分类模型", "mutil_train/train/classification_model.py (Focust/SimpleCNN + CfC)"),
        ],
    )

    doc.add_page_break()

    doc.add_heading("摘要", level=1)
    doc.add_paragraph(
        "面向含食品残渣干扰的培养皿时序菌落生长图像，本系统提出一种“传统机器视觉预处理（HCP）+ 时序深度学习分类”的三层混合架构。"
        "第一层 HCP 通过鲁棒背景建模与变化检测在全帧上快速提出候选区域；第二层二分类网络利用时序特征区分真实菌落与残渣等伪目标；"
        "第三层多分类网络在确认菌落上完成菌种识别。该方案兼顾泛化与轻量化：以物理先验降低误检，以 ROI 级别时序网络提升判别与分类稳定性，"
        "最终实现对时序图像序列的菌落计数与分类输出。"
    )
    doc.add_paragraph("关键词：HCP；背景建模；时序分类；CfC；NCP；二分类；多分类；菌落计数；食品残渣干扰")

    _section_repo_basis(doc)
    _section_problem_background(doc)
    _section_system_overview(doc, cfg)
    _section_hcp_algorithm(doc)
    _section_binary_model(doc)
    _section_multiclass_model(doc, cfg)
    _section_hpo_optuna(doc)
    _section_inference_and_counting(doc, cfg)
    _section_comparison_and_hybrid_rationale(doc)
    _section_limitations_and_open_issues(doc)
    _section_binary_dataset_construction(doc)
    _section_multiclass_dataset_prelabel_and_correction(doc)
    _section_appendix_configs(doc, cfg)

    out_dir = REPO_ROOT / "reports"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "Focust系统_HCP_二分类初筛_多分类复筛_时序菌落计数与分类_技术白皮书.docx"
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    path = main()
    print(str(path))
